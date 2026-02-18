"""
Phase 5 — AI Agent for parsing maintenance manuals (PDF/DOCX).

Isolated module: downloads file from Bitrix24 Disk → extracts text →
sends to LLM (OpenAI / Claude / Gemini / Grok) → returns structured
maintenance intervals + tasks.

Does NOT import any core SCADA modules (no models, no Redis, no WebSocket).
"""
import io
import json
import logging
from typing import Optional

import httpx

from config import settings

logger = logging.getLogger("scada.ai_agent")

# ---------------------------------------------------------------------------
# System prompt for maintenance manual parsing (Russian DGU/GPU context)
# ---------------------------------------------------------------------------
SYSTEM_PROMPT = """Ты — AI-агент для парсинга регламентов технического обслуживания (ТО) дизельных и газопоршневых генераторных установок (ДГУ/ГПУ).

Твоя задача: из текста технического мануала извлечь структурированный регламент ТО.

ФОРМАТ ОТВЕТА (строго JSON):
{
  "name": "Название регламента (из документа или 'Регламент ТО')",
  "description": "Краткое описание источника",
  "intervals": [
    {
      "code": "to1",
      "name": "ТО-1",
      "hours": 250,
      "sort_order": 0,
      "tasks": [
        {"text": "Описание работы", "is_critical": true, "sort_order": 0}
      ]
    }
  ]
}

ПРАВИЛА ИЗВЛЕЧЕНИЯ:
1. Интервалы: найди все уровни ТО (ТО-1, ТО-2, ТО-3, ТО-4 и т.д.) с их интервалами в моточасах.
2. Если интервалы не пронумерованы явно — пронумеруй по возрастанию часов (to1, to2...).
3. Типичные интервалы ДГУ: 250ч, 500ч, 1000ч, 2000ч, 4000ч, 8000ч — но бери из документа.
4. Для каждого интервала извлеки список работ/задач.
5. is_critical=true для: замена масла, замена фильтров, проверки безопасности, замена ремней ГРМ, работы влияющие на ресурс двигателя.
6. is_critical=false для: визуальный осмотр, проверка уровней, очистка, обновление журнала.
7. Текст задач — краткий, на русском, в формате действия ("Замена...", "Проверка...", "Регулировка...").
8. sort_order — по порядку появления в документе, начиная с 0.
9. code — латинские обозначения: to1, to2, to3 и т.д.
10. Если в документе указано "включая все работы предыдущего ТО" — добавь задачу "Все работы предыдущего ТО" с is_critical=true.
11. Если документ не содержит регламент ТО — верни пустой intervals: [] и описание причины в description.

Отвечай ТОЛЬКО валидным JSON без markdown, без комментариев, без пояснений."""

# Max characters to send to LLM
MAX_TEXT_CHARS = 80_000


class AIAgentError(Exception):
    """Base exception for AI agent operations."""
    pass


class MaintenanceDocumentParser:
    """
    Parses maintenance manuals from Bitrix24 Disk using LLM.
    Supports: OpenAI, Claude, Gemini, Grok.
    Completely isolated from core SCADA — no DB, no Redis, no WebSocket.
    """

    def __init__(
        self,
        provider: str = "openai",
        api_key: str = "",
        model: str = "",
    ):
        self.provider = provider
        self.api_key = api_key
        self.timeout = settings.AI_TIMEOUT

        if not self.api_key:
            raise AIAgentError(f"API ключ для {provider} не настроен")

        # Set default models per provider
        if not model:
            model = {
                "openai": settings.OPENAI_MODEL,
                "claude": settings.CLAUDE_MODEL,
                "gemini": settings.GEMINI_MODEL,
                "grok": settings.GROK_MODEL,
            }.get(provider, "gpt-4o")
        self.model = model

        # Initialize provider-specific client (only for OpenAI/Grok SDK)
        if provider in ("openai", "grok"):
            from openai import AsyncOpenAI
            base_url = "https://api.x.ai/v1" if provider == "grok" else None
            self.client = AsyncOpenAI(
                api_key=self.api_key,
                timeout=self.timeout,
                base_url=base_url,
            )

    # ------------------------------------------------------------------
    # Step 1: Download file from Bitrix24 Disk
    # ------------------------------------------------------------------
    async def download_file_from_bitrix(
        self, webhook_url: str, file_id: int
    ) -> tuple[bytes, str]:
        """
        Download file from Bitrix24 Disk.
        1. disk.file.get → get metadata (NAME, DOWNLOAD_URL)
        2. Download binary content from DOWNLOAD_URL
        Returns (file_bytes, filename).
        """
        webhook_url = webhook_url.rstrip("/")

        async with httpx.AsyncClient(timeout=30, verify=False) as http:
            # Get file metadata
            logger.info("Fetching file metadata from Bitrix24: file_id=%d", file_id)
            resp = await http.post(
                f"{webhook_url}/disk.file.get",
                json={"id": file_id},
            )
            resp.raise_for_status()
            data = resp.json()

            result = data.get("result")
            if not result:
                error = data.get("error_description", data.get("error", "Unknown"))
                raise AIAgentError(f"Битрикс24 вернул ошибку: {error}")

            filename = result.get("NAME", f"file_{file_id}")
            download_url = result.get("DOWNLOAD_URL")

            if not download_url:
                raise AIAgentError(
                    f"Файл {filename} не имеет DOWNLOAD_URL. "
                    "Проверьте права доступа webhook (disk scope)."
                )

            # Download file content
            logger.info("Downloading file: %s (%s)", filename, download_url[:80])
            file_resp = await http.get(download_url)
            file_resp.raise_for_status()

            file_bytes = file_resp.content
            logger.info(
                "Downloaded %s: %d bytes", filename, len(file_bytes)
            )
            return file_bytes, filename

    # ------------------------------------------------------------------
    # Step 2: Extract text from PDF or DOCX
    # ------------------------------------------------------------------
    def extract_text(self, file_bytes: bytes, filename: str) -> str:
        """Extract text from PDF or DOCX file."""
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if ext == "pdf":
            return self._extract_pdf(file_bytes)
        elif ext in ("docx", "doc"):
            return self._extract_docx(file_bytes)
        else:
            raise AIAgentError(
                f"Формат .{ext} не поддерживается. Используйте PDF или DOCX."
            )

    def _extract_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF using pypdf."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise AIAgentError("pypdf не установлен")

        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)

        full_text = "\n\n".join(pages)
        if not full_text.strip():
            raise AIAgentError(
                "PDF не содержит извлекаемого текста. "
                "Возможно, это скан без OCR."
            )
        logger.info("PDF: extracted %d chars from %d pages", len(full_text), len(reader.pages))
        return full_text

    def _extract_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX using python-docx (including tables)."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise AIAgentError("python-docx не установлен")

        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables (important — maintenance manuals often use tables)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        full_text = "\n".join(parts)
        if not full_text.strip():
            raise AIAgentError("DOCX не содержит текста.")
        logger.info("DOCX: extracted %d chars", len(full_text))
        return full_text

    # ------------------------------------------------------------------
    # Step 3: Parse document text with LLM
    # ------------------------------------------------------------------
    async def _call_openai(self, text: str, filename: str) -> str:
        """Call OpenAI or Grok (OpenAI-compatible) API."""
        from openai import APITimeoutError, APIError

        try:
            response = await self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {
                        "role": "user",
                        "content": f"Документ: {filename}\n\nТекст документа:\n{text}",
                    },
                ],
                response_format={"type": "json_object"},
                temperature=0.1,
            )
        except APITimeoutError:
            raise AIAgentError(
                f"Таймаут {self.provider} ({self.timeout}с). "
                "Попробуйте документ меньшего размера."
            )
        except APIError as e:
            raise AIAgentError(f"Ошибка {self.provider} API: {e.message}")

        return response.choices[0].message.content

    async def _call_claude(self, text: str, filename: str) -> str:
        """Call Anthropic Claude API via httpx."""
        async with httpx.AsyncClient(timeout=self.timeout) as http:
            try:
                resp = await http.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": self.api_key,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                    },
                    json={
                        "model": self.model,
                        "max_tokens": 8192,
                        "system": SYSTEM_PROMPT,
                        "messages": [
                            {
                                "role": "user",
                                "content": f"Документ: {filename}\n\nТекст документа:\n{text}",
                            },
                        ],
                        "temperature": 0.1,
                    },
                )
            except httpx.TimeoutException:
                raise AIAgentError(
                    f"Таймаут Claude ({self.timeout}с). "
                    "Попробуйте документ меньшего размера."
                )

            if resp.status_code != 200:
                err = resp.json().get("error", {}).get("message", resp.text)
                raise AIAgentError(f"Ошибка Claude API: {err}")

            data = resp.json()
            return data.get("content", [{}])[0].get("text", "")

    async def _call_gemini(self, text: str, filename: str) -> str:
        """Call Google Gemini API via httpx."""
        url = (
            f"https://generativelanguage.googleapis.com/v1beta/models/"
            f"{self.model}:generateContent?key={self.api_key}"
        )
        prompt = f"{SYSTEM_PROMPT}\n\n---\n\nДокумент: {filename}\n\nТекст документа:\n{text}"

        async with httpx.AsyncClient(timeout=self.timeout) as http:
            try:
                resp = await http.post(
                    url,
                    json={
                        "contents": [{"parts": [{"text": prompt}]}],
                        "generationConfig": {
                            "responseMimeType": "application/json",
                            "temperature": 0.1,
                            "maxOutputTokens": 8192,
                        },
                    },
                )
            except httpx.TimeoutException:
                raise AIAgentError(
                    f"Таймаут Gemini ({self.timeout}с). "
                    "Попробуйте документ меньшего размера."
                )

            if resp.status_code != 200:
                err = resp.json().get("error", {}).get("message", resp.text)
                raise AIAgentError(f"Ошибка Gemini API: {err}")

            data = resp.json()
            return (
                data.get("candidates", [{}])[0]
                .get("content", {})
                .get("parts", [{}])[0]
                .get("text", "")
            )

    async def parse_document(self, file_bytes: bytes, filename: str) -> dict:
        """
        Main orchestration: extract text → truncate → LLM → JSON.
        Returns structured maintenance template dict.
        """
        # Extract text
        text = self.extract_text(file_bytes, filename)

        # Truncate if too long
        if len(text) > MAX_TEXT_CHARS:
            logger.warning(
                "Text too long (%d chars), truncating to %d",
                len(text), MAX_TEXT_CHARS,
            )
            text = text[:MAX_TEXT_CHARS] + "\n\n[...текст обрезан...]"

        # Call LLM based on provider
        logger.info(
            "Sending %d chars to %s (%s) for parsing...",
            len(text), self.provider, self.model,
        )

        if self.provider in ("openai", "grok"):
            content = await self._call_openai(text, filename)
        elif self.provider == "claude":
            content = await self._call_claude(text, filename)
        elif self.provider == "gemini":
            content = await self._call_gemini(text, filename)
        else:
            raise AIAgentError(f"Неизвестный провайдер: {self.provider}")

        logger.info("%s response: %d chars", self.provider, len(content or ""))

        # Parse response — strip markdown code fences if present
        if content:
            content = content.strip()
            if content.startswith("```"):
                # Remove ```json ... ``` wrapping
                lines = content.split("\n")
                if lines[0].startswith("```"):
                    lines = lines[1:]
                if lines and lines[-1].strip() == "```":
                    lines = lines[:-1]
                content = "\n".join(lines)

        try:
            result = json.loads(content)
        except json.JSONDecodeError:
            raise AIAgentError(
                f"{self.provider} вернул некорректный JSON. Повторите попытку."
            )

        # Validate structure
        if "intervals" not in result:
            result["intervals"] = []
        if "name" not in result:
            result["name"] = f"Регламент ТО из {filename}"
        if "description" not in result:
            result["description"] = f"Извлечено из: {filename}"

        logger.info(
            "Parsed %d intervals from %s via %s",
            len(result["intervals"]), filename, self.provider,
        )
        return result

    # ------------------------------------------------------------------
    # Step 4: End-to-end: download from Bitrix + parse
    # ------------------------------------------------------------------
    async def parse_bitrix_file(
        self, webhook_url: str, file_id: int, filename: Optional[str] = None
    ) -> dict:
        """
        End-to-end: download file from Bitrix24 → extract text → LLM → JSON.
        """
        file_bytes, bx_filename = await self.download_file_from_bitrix(
            webhook_url, file_id
        )
        # Use provided filename or the one from Bitrix
        actual_filename = filename or bx_filename

        result = await self.parse_document(file_bytes, actual_filename)

        # Add preview of raw text for debug
        text = self.extract_text(file_bytes, actual_filename)
        result["raw_text_preview"] = text[:500]

        return result
